#!/usr/bin/env python3
"""Generate IntelliRAG Test Case Specification v1.0 as a .docx file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_shading(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def cell_txt(cell, text, bold=False, sz=9, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.font.name = "Arial"; r.font.size = Pt(sz); r.bold = bold
    if color: r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def hdr_row(tbl, hdrs):
    for i, h in enumerate(hdrs):
        set_shading(tbl.rows[0].cells[i], "D6E4F0")
        cell_txt(tbl.rows[0].cells[i], h, bold=True, color="003366")

def add_tc(doc, tc):
    fields = [("Test ID",tc[0]),("Test Name",tc[1]),("Category",tc[2]),
              ("Linked FR",tc[3]),("Priority",tc[4]),("Preconditions",tc[5]),
              ("Test Steps",tc[6]),("Test Data",tc[7]),("Expected Result",tc[8]),
              ("Postconditions",tc[9])]
    t = doc.add_table(rows=10, cols=2)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,(l,v) in enumerate(fields):
        set_shading(t.rows[i].cells[0], "D6E4F0")
        cell_txt(t.rows[i].cells[0], l, bold=True, color="003366")
        cell_txt(t.rows[i].cells[1], v)
        t.rows[i].cells[0].width = Cm(3.5)
        t.rows[i].cells[1].width = Cm(13)
    doc.add_paragraph("")

# TC format: (id, name, cat, fr, pri, precond, steps, data, expected, postcond)
TCS = []
def t(id,nm,cat,fr,pri,pre,stp,dat,exp,post):
    TCS.append((id,nm,cat,fr,pri,pre,stp,dat,exp,post))

# ===== MODULE A: FILE INGESTION =====
# FR-001
t("TC-FR001-01","Upload single PDF via file picker","Happy Path","FR-001","P1",
"User logged in as Analyst. No duplicate file.",
"1. Navigate to Document Upload.\n2. Click Browse Files.\n3. Select Q4_Report.pdf (2.4 MB).\n4. Wait for upload.",
"File: Q4_Report.pdf, 2.4 MB, PDF",
"Progress bar 0-100%. File in library with Pending (gray) badge. SHA-256 checksum computed. Audit log document.upload created.",
"Document record: ingestion_status=pending, gcs_uri populated.")

t("TC-FR001-02","Upload all 14 supported formats via drag-and-drop","Happy Path","FR-001","P1",
"User logged in as Data Manager.",
"1. Drag 14 files (PDF,DOCX,DOC,XLSX,XLS,CSV,MD,JPEG,PNG,TIFF,BMP,GIF,WEBP,PDF) onto drop zone.\n2. Observe progress.",
"14 files, one per format, each <5 MB",
"All 14 accepted. Individual progress bars. All appear with Pending status. 14 audit entries.",
"14 document records in DB.")

t("TC-FR001-03","Reject unsupported file format (.exe)","Negative","FR-001","P1",
"User logged in as Analyst.",
"1. Click Browse Files.\n2. Select malware.exe.\n3. Attempt upload.",
"File: malware.exe, 1 MB",
"Error: 'Unsupported file format. Supported formats: PDF, DOCX, DOC, XLSX, XLS, CSV, MD, JPEG, PNG, TIFF, BMP, GIF, WEBP.' No upload occurs.",
"No document record. No file in GCS.")

t("TC-FR001-04","Reject file exceeding 100 MB","Boundary","FR-001","P1",
"User logged in as Analyst.",
"1. Select large_dataset.pdf (101 MB).\n2. Attempt upload.",
"File: large_dataset.pdf, 104,857,601 bytes",
"Error: 'File exceeds maximum size of 100 MB.' File not uploaded.",
"No document record in DB.")

t("TC-FR001-05","Accept file exactly at 100 MB boundary","Boundary","FR-001","P2",
"User logged in as Analyst.",
"1. Select boundary.pdf (exactly 104,857,600 bytes).\n2. Upload.",
"File: boundary.pdf, 104,857,600 bytes",
"Upload succeeds. File in library with Pending status.",
"Document record: file_size_bytes=104857600.")

t("TC-FR001-06","Duplicate SHA-256 checksum detection","Happy Path","FR-001","P1",
"Q4_Report.pdf already uploaded.",
"1. Upload Q4_Report_copy.pdf (identical content).\n2. Observe prompt.\n3. Click Upload anyway.",
"File identical SHA-256 to existing Q4_Report.pdf",
"Prompt: 'A file with identical content already exists (Q4_Report.pdf). Upload anyway?' After confirm, new version uploaded.",
"Two records with same checksum_sha256.")

t("TC-FR001-07","Reject batch exceeding 20 files","Boundary","FR-001","P2",
"User logged in as Analyst.",
"1. Select 21 PDF files via picker.\n2. Attempt upload.",
"21 files: file_01.pdf to file_21.pdf, 500 KB each",
"Error indicating max 20 files per batch. 21st file rejected.",
"Max 20 files uploaded.")

t("TC-FR001-08","File name sanitization","Happy Path","FR-001","P2",
"User logged in.",
"1. Upload 'Report @#$% (Final) v2!.pdf'.\n2. Check stored filename.",
"Filename with special chars, 1 MB",
"Upload succeeds. stored_filename has special chars removed (except hyphens/underscores). Max 255 chars.",
"Safe stored_filename in DB. Original preserved in original_filename.")

t("TC-FR001-09","Viewer role blocked from uploading","Permission","FR-001","P1",
"Logged in as Viewer.",
"1. Navigate to Upload page.\n2. Attempt upload or call POST /api/v1/documents/upload.",
"User role: Viewer",
"Upload denied. UI: button disabled/hidden. API: HTTP 403 FORBIDDEN.",
"No document created.")

t("TC-FR001-10","Network retry with exponential backoff","Negative","FR-001","P2",
"Network unstable during upload.",
"1. Begin uploading large_report.pdf (50 MB).\n2. Simulate network drop at 40%.\n3. Observe retries.",
"File: 50 MB. Network drops at 40%.",
"Auto-retry up to 3 times (2s, 4s, 8s backoff). On success: upload completes. On failure: 'Upload failed due to network error. Please try again.'",
"On success: record created. On failure: no partial record.")

# FR-002
t("TC-FR002-01","XLSX converted to CSV within 30 seconds","Happy Path","FR-002","P1",
"User logged in. LibreOffice available.",
"1. Upload financial_data.xlsx (3 MB, single sheet, 500 rows).\n2. Wait for conversion.\n3. Check Ingestion Jobs.",
"financial_data.xlsx, 3 MB, 1 sheet",
"CSV created within 30s. Job: job_type=convert, status=completed. CSV in GCS at artifacts/{doc_id}/csv/. Original preserved.",
"Ingestion job completed. CSV artifact in GCS.")

t("TC-FR002-02","Multi-sheet XLSX produces separate CSVs","Happy Path","FR-002","P1",
"User logged in.",
"1. Upload quarterly_data.xlsx (4 sheets: Q1-Q4).\n2. Wait for conversion.\n3. Verify 4 CSV artifacts.",
"quarterly_data.xlsx, 4 sheets, 200 rows each",
"4 CSV files created, one per sheet, all linked to parent doc. Completes within 30s.",
"4 CSV artifacts in GCS.")

t("TC-FR002-03","DOC converted to DOCX via LibreOffice","Happy Path","FR-002","P2",
"User logged in.",
"1. Upload legacy_report.doc (2 MB, 15 pages).\n2. Wait.\n3. Verify DOCX artifact.",
"legacy_report.doc, 2 MB",
"DOC converted to DOCX. Artifact in GCS. Job status=completed. Original DOC preserved.",
"DOCX artifact exists.")

t("TC-FR002-04","Conversion failure after 3 retries sets status to failed","Negative","FR-002","P1",
"Uploaded a corrupted XLSX.",
"1. Upload corrupt_data.xlsx.\n2. Wait for 3 retry attempts.\n3. Check status.",
"corrupt_data.xlsx, internally corrupted",
"Fails 3 times. Status=failed. Notification: 'corrupt_data.xlsx failed to process: [error]. Please review and retry.' Original file preserved.",
"ingestion_status=failed. attempt_count=3.")

t("TC-FR002-05","Markdown converted to plain text preserving structure","Happy Path","FR-002","P2",
"User logged in.",
"1. Upload readme.md with H1-H3, bullets, code blocks.\n2. Wait.\n3. Inspect plain text.",
"readme.md, 50 KB, markdown with headers/lists/code",
"Plain text with headers and structure preserved. Conversion completed.",
"Plain text artifact in GCS.")

t("TC-FR002-06","PDF table extraction as JSON","Happy Path","FR-002","P1",
"PDF with embedded tables.",
"1. Upload sales_report.pdf (3 tables).\n2. Wait.\n3. Verify JSON output.",
"sales_report.pdf with Region/Revenue/Growth tables",
"Tables as JSON: {\"headers\":[\"Region\",\"Revenue\",\"Growth\"],\"rows\":[[\"APAC\",\"$1.2B\",\"15%\"],...]}. chunk_type=table.",
"3 table chunks with valid JSON.")

# FR-003
t("TC-FR003-01","JPEG automatically OCR-processed","Happy Path","FR-003","P1",
"Document AI available.",
"1. Upload scanned_invoice.jpeg (300 DPI).\n2. Wait for OCR.\n3. Verify text.",
"scanned_invoice.jpeg, 800 KB, 300 DPI, text: 'Invoice #1234, $5,000.00'",
"OCR applied. ocr_applied=true. Text: 'Invoice #1234', '$5,000.00'. Accuracy >=95%.",
"ocr_applied=true. Extracted text stored.")

t("TC-FR003-02","PDF <10 chars/page triggers OCR","Happy Path","FR-003","P1",
"Scanned PDF uploaded.",
"1. Upload scanned_contract.pdf (10 pages, <10 chars/page extractable).\n2. Verify OCR triggered.",
"scanned_contract.pdf, 8 MB, 10 pages",
"<10 chars/page detected. OCR invoked for all pages. ocr_applied=true. Structure preserved.",
"Full text extracted. OCR applied.")

t("TC-FR003-03","Native text PDF skips OCR","Negative","FR-003","P2",
"PDF has extractable text.",
"1. Upload digital_report.pdf (500+ chars/page).\n2. Check ocr_applied.",
"digital_report.pdf, 2 MB, 20 pages, >500 chars/page",
"OCR NOT invoked. ocr_applied=false. Text extracted directly.",
"ocr_applied=false.")

t("TC-FR003-04","OCR page timeout handled gracefully","Negative","FR-003","P2",
"One page causes OCR timeout.",
"1. Upload complex_scan.pdf (20 pages, page 15 times out).\n2. Check results.",
"complex_scan.pdf, 20 pages. Page 15 exceeds 120s timeout.",
"Pages 1-14 and 16-20 processed. Page 15 skipped and logged. Processing continues.",
"19/20 pages extracted. Timeout logged.")

t("TC-FR003-05","OCR confidence <0.7 triggers warning","Boundary","FR-003","P2",
"Low quality scan.",
"1. Upload blurry_scan.jpeg (72 DPI, poor quality).\n2. Check ingestion log.",
"blurry_scan.jpeg, 200 KB, 72 DPI",
"OCR processes image. Confidence <0.7. Warning in log: 'Low OCR confidence (0.45) for page 1'.",
"Warning logged. ocr_applied=true.")

t("TC-FR003-06","Parallel OCR for multi-page scans","Happy Path","FR-003","P2",
"Document AI supports parallelism.",
"1. Upload large_scan.pdf (30 scanned pages).\n2. Monitor processing time.",
"large_scan.pdf, 25 MB, 30 pages",
"Parallel processing (up to 10 concurrent pages). All 30 pages extracted. Time << sequential.",
"All pages extracted with parallelism.")

# FR-004
t("TC-FR004-01","Tables converted to JSON with headers/rows","Happy Path","FR-004","P1",
"Tables detected in document.",
"1. Upload report_with_tables.pdf.\n2. Wait.\n3. Inspect chunks.",
"Table: Region|Revenue|Growth / APAC|$1.2B|15%",
"JSON: {\"headers\":[\"Region\",\"Revenue\",\"Growth\"],\"rows\":[[\"APAC\",\"$1.2B\",\"15%\"]]}. chunk_type=table.",
"Chunk with valid JSON structure.")

t("TC-FR004-02","Nested lists preserve 5 levels of hierarchy","Happy Path","FR-004","P2",
"Document with deep nesting.",
"1. Upload outline.docx (5-level nested list).\n2. Verify chunk content.",
"5-level nested list: 1 > 1.1 > 1.1.1 > 1.1.1.1 > 1.1.1.1.1",
"All 5 levels preserved with proper indentation. No loss of items.",
"Chunks preserve hierarchy.")

t("TC-FR004-03","Unicode and special characters handled","Boundary","FR-004","P2",
"Document with unicode/ligatures.",
"1. Upload international.pdf with accents, ligatures, em-dashes.\n2. Verify output.",
"Text with accented chars, fi/fl ligatures, em-dashes, curly quotes",
"All special chars preserved. Ligatures resolved. No garbled text.",
"Correct unicode in chunks.")

# ===== MODULE B: STORAGE, CHUNKING, VECTORIZATION =====
# FR-005
t("TC-FR005-01","File stored at correct GCS path","Happy Path","FR-005","P1",
"GCS bucket exists.",
"1. Upload Q4_Report.pdf on 2026-03-17.\n2. Verify GCS path.\n3. Check gcs_uri.",
"Upload date: 2026-03-17, doc_id: doc-abc-123",
"Path: gs://[bucket]/documents/2026/03/doc-abc-123/Q4_Report.pdf. gcs_uri populated. AES-256 encrypted.",
"GCS URI in DB. File in GCS.")

t("TC-FR005-02","Artifact stored at correct path","Happy Path","FR-005","P2",
"XLSX converted to CSV.",
"1. Upload data.xlsx (doc_id: doc-xyz-789).\n2. Verify artifact path.",
"doc_id: doc-xyz-789",
"CSV at: gs://[bucket]/artifacts/doc-xyz-789/csv/data.csv. Original XLSX preserved.",
"Both files in GCS.")

t("TC-FR005-03","Storage lifecycle policy: 90d Nearline, 365d Coldline","Integration","FR-005","P3",
"Bucket configured.",
"1. Verify lifecycle config.\n2. Check 90-day Nearline rule.\n3. Check 365-day Coldline rule.",
"Lifecycle: 90d -> Nearline, 365d -> Coldline",
"Policies configured on bucket. Files transition correctly. Remain accessible.",
"Lifecycle rules active.")

# FR-006
t("TC-FR006-01","Fixed-size chunking: 512 tokens, 50 overlap","Happy Path","FR-006","P1",
"Document text extracted. Strategy: fixed_size.",
"1. Upload long_document.pdf (10,000 tokens).\n2. Apply fixed_size (512/50).\n3. Verify chunks.",
"10,000 tokens. Strategy: fixed_size, size=512, overlap=50",
"~22 chunks. Each ~512 tokens. Adjacent overlap 50 tokens. chunking_strategy=fixed_size.",
"Chunks in DB with correct counts.")

t("TC-FR006-02","Semantic chunking by topic shifts","Happy Path","FR-006","P1",
"Strategy: semantic.",
"1. Upload article.pdf (3 sections).\n2. Apply semantic chunking.\n3. Inspect boundaries.",
"3 sections (~1000 tokens each)",
"Chunks align with topic shifts/sentence boundaries. No mid-sentence splits.",
"Topic-coherent chunks.")

t("TC-FR006-03","Table-aware chunking keeps tables intact","Happy Path","FR-006","P1",
"Document has tables. Strategy: table_aware.",
"1. Upload report.pdf (2 tables, 30 rows each).\n2. Apply table_aware.",
"2 tables, 30 rows each (<50 threshold)",
"Tables as single chunks. Surrounding text chunked separately. chunk_type=table.",
"Table chunks intact.")

t("TC-FR006-04","Chunk >10,000 chars split at sentence boundary","Boundary","FR-006","P1",
"Extremely long paragraph.",
"1. Upload doc with 15,000-char paragraph.\n2. Apply paragraph chunking.\n3. Verify split.",
"15,000-char continuous paragraph",
"Split into 2 chunks at nearest sentence boundary. No chunk >10,000 chars.",
"All chunks <=10,000 chars.")

t("TC-FR006-05","Chunk <50 chars merged with adjacent","Boundary","FR-006","P2",
"Very short paragraphs.",
"1. Upload doc: 'OK.' (3 chars), 'Yes.' (4 chars), normal 500-char para.\n2. Apply paragraph chunking.",
"Paragraphs: 3 chars, 4 chars, 500 chars",
"Short paras merged with adjacent. No chunk <50 chars in output.",
"No chunks below 50 chars.")

t("TC-FR006-06","Large table >50 rows split into 25-row sub-tables","Boundary","FR-006","P2",
"Table with 75 rows.",
"1. Upload doc with 75-row table.\n2. Apply table_aware chunking.",
"75-row table, 5 columns",
"3 sub-tables of 25 rows. Header repeated in each. chunk_type=table.",
"3 table chunks with headers.")

# FR-007
t("TC-FR007-01","PDF metadata extracted","Happy Path","FR-007","P1",
"PDF with embedded metadata.",
"1. Upload report.pdf (Author=John Doe, Title=Q4 Report).\n2. Check Documents.metadata.",
"PDF metadata: Author=John Doe, Title=Q4 Report, Keywords=finance,quarterly",
"metadata JSONB: author=John Doe, title=Q4 Report, keywords=[finance,quarterly]. Indexed for search.",
"Metadata populated and searchable.")

t("TC-FR007-02","Add custom tags (up to 20)","Happy Path","FR-007","P2",
"Document exists.",
"1. Add tags: finance, quarterly, apac-region.\n2. Save.",
"Tags: finance, quarterly, apac-region",
"Tags saved. Searchable. Displayed on detail page.",
"Tags in metadata.")

t("TC-FR007-03","Reject tag >50 characters","Boundary","FR-007","P2",
"Document exists.",
"1. Add tag of 51 characters.\n2. Save.",
"Tag: 51-char string",
"Validation error: max 50 chars. Tag not saved.",
"No invalid tag stored.")

t("TC-FR007-04","Reject 21st custom tag","Boundary","FR-007","P2",
"Document has 20 tags.",
"1. Attempt to add 21st tag.\n2. Save.",
"20 existing tags. New: extra-tag",
"Error: max 20 tags. 21st rejected.",
"Still exactly 20 tags.")

t("TC-FR007-05","Language detection stores ISO 639-1 code","Happy Path","FR-007","P2",
"English document uploaded.",
"1. Upload english_report.pdf.\n2. Check Documents.language.",
"English text document",
"language='en' (ISO 639-1). Stored and searchable.",
"Documents.language=en.")

# FR-008
t("TC-FR008-01","Embeddings generated within 60s","Happy Path","FR-008","P1",
"Chunks created. Vertex AI available.",
"1. Upload report.pdf (50 chunks).\n2. Monitor embedding time.\n3. Verify vectors.",
"50 chunks, Vertex AI text-embedding-004",
"50 embeddings in <60s. VECTOR(768) stored. Status advances to vectorized. HNSW index covers them.",
"All 50 chunks have 768-dim embeddings.")

t("TC-FR008-02","Batch embedding: 1000 chunks","Happy Path","FR-008","P2",
"Large document producing 1000 chunks.",
"1. Upload very large doc (1000 chunks).\n2. Verify batch API.",
"1000 chunks",
"All 1000 processed in single batch call. 768-dim embeddings stored.",
"1000 embeddings in DB.")

t("TC-FR008-03","Embedding failure retries up to 3 times","Negative","FR-008","P1",
"Vertex AI transient errors.",
"1. Upload doc (10 chunks).\n2. Simulate API failure.\n3. Observe retries.",
"Vertex AI 503 on first 2 attempts",
"Retries up to 3 times. On success: embeddings created. On failure: status remains chunked.",
"Embeddings generated after retry.")

# ===== MODULE C: KNOWLEDGE GRAPH =====
# FR-009
t("TC-FR009-01","Entities extracted with confidence >=0.75","Happy Path","FR-009","P1",
"Chunks ready. LangChain configured.",
"1. Process chunks: 'Rajesh Kumar, CEO of ADS Softek, Q4 APAC revenue $1.2B.'\n2. Verify entities.",
"Text with Person, Organization, Location, Amount",
"Entities: Rajesh Kumar (Person,0.92), ADS Softek (Org,0.97), APAC (Location,0.95), $1.2B (Amount,0.88). All >=0.75. source_ids populated. >=100 chunks/min.",
"KG nodes with correct types and scores.")

t("TC-FR009-02","Entities <0.75 confidence discarded","Boundary","FR-009","P1",
"Ambiguous entity in text.",
"1. Process chunk with ambiguous mention.\n2. Check if low-confidence entity stored.",
"Ambiguous reference: 'Project Alpha mentioned in passing' (confidence 0.60)",
"Entity with confidence 0.60 discarded. No KG node created for it.",
"No nodes with confidence <0.75.")

t("TC-FR009-03","Entity deduplication >=90% Levenshtein similarity","Happy Path","FR-009","P2",
"Same entity in multiple chunks with name variants.",
"1. Process 'ADS Softek', 'ADS Softek Pvt Ltd', 'ADS'.\n2. Check dedup.",
"Name variants across 3 chunks",
"Single canonical entity 'ADS Softek' with aliases. occurrence_count aggregated. source_chunk_ids from all chunks.",
"Single entity with aliases.")

# FR-010
t("TC-FR010-01","Relationships extracted between entities","Happy Path","FR-010","P1",
"Entities extracted from document.",
"1. Process entities.\n2. Verify edges.\n3. Check types and weights.",
"Entities: Rajesh Kumar, ADS Softek, APAC Region",
"Edges: WORKS_AT (weight 0.97), OPERATES_IN (weight 0.95). source_chunk_ids populated. Bi-directional query works.",
"KG edges with types and weights.")

t("TC-FR010-02","Duplicate relationships merged, weight=max","Happy Path","FR-010","P2",
"Same relationship from 2 documents.",
"1. Process Doc A: WORKS_AT (weight 0.85).\n2. Process Doc B: same (weight 0.97).\n3. Check edge.",
"Same relationship from 2 docs, weights 0.85 and 0.97",
"Single edge. Weight updated to 0.97. source_chunk_ids includes both docs.",
"Single edge, weight=0.97.")

t("TC-FR010-03","Bi-directional relationship query","Happy Path","FR-010","P2",
"KG populated.",
"1. Query entities related TO ADS Softek.\n2. Query entities FROM ADS Softek.",
"ADS Softek with incoming/outgoing edges",
"TO: Rajesh Kumar (WORKS_AT). FROM: APAC Region (OPERATES_IN). Both directions work.",
"Bi-directional queries correct.")

# FR-011
t("TC-FR011-01","3-hop graph traversal under 500ms","Happy Path","FR-011","P1",
"KG with 10,000+ nodes.",
"1. Execute 3-hop query from ADS Softek.\n2. Measure time.",
"KG: 10,000 nodes, 25,000 edges",
"Completes <500ms. Returns multi-hop entities via recursive CTE.",
"Results in <500ms.")

t("TC-FR011-02","Admin triggers full KG rebuild","Happy Path","FR-011","P1",
"Admin logged in. KG exists.",
"1. Admin > Knowledge Graph > Rebuild.\n2. Confirm.\n3. Wait.",
"KG: 5,000 nodes, 12,000 edges",
"Rebuild completes. Notification: 'KG rebuild completed. [X] nodes, [Y] edges.' Stats updated.",
"KG rebuilt. Stats reflect new counts.")

t("TC-FR011-03","Analyst cannot trigger KG rebuild","Permission","FR-011","P1",
"Logged in as Analyst.",
"1. Attempt Admin > KG tab.\n2. API: POST /api/v1/admin/knowledge-graph/rebuild.",
"Role: Analyst",
"Tab not visible. API: HTTP 403 FORBIDDEN.",
"No rebuild initiated.")

# ===== MODULE D: RAG SEARCH =====
# FR-012
t("TC-FR012-01","Step-back expanded query generated","Happy Path","FR-012","P1",
"Active conversation. KB populated.",
"1. Ask: 'What was APAC Q4 revenue?'\n2. Check expanded_intent.",
"Question: 'What was APAC Q4 revenue?'",
"Expanded: broader intent about APAC financial performance. Stored in Messages.expanded_intent. Generated <500ms. Both used for retrieval.",
"expanded_intent populated.")

t("TC-FR012-02","Short query (<3 words) skips step-back","Boundary","FR-012","P1",
"Active conversation.",
"1. Ask: 'APAC revenue' (2 words).\n2. Check expanded_intent.",
"Question: 'APAC revenue'",
"Step-back skipped. expanded_intent null or same as original. Query used as-is.",
"expanded_intent null.")

t("TC-FR012-03","Step-back completes within 500ms","Happy Path","FR-012","P2",
"LLM available.",
"1. Ask complex question.\n2. Measure expansion time.",
"Long question about employee satisfaction",
"Expanded intent generated <500ms.",
"Latency within threshold.")

# FR-013
t("TC-FR013-01","KG entities matched, 2-hop retrieval with NL context","Happy Path","FR-013","P1",
"KG populated. Query mentions known entities.",
"1. Ask: 'What is ADS Softek revenue in APAC?'\n2. Verify KG match and 2-hop retrieval.",
"'ADS Softek' and 'APAC' in KG",
"Entities matched. 2-hop related retrieved. Context: 'ADS Softek operates in APAC Region.' Graph retrieval <300ms. retrieved_node_ids stored.",
"Graph context generated. Node IDs recorded.")

t("TC-FR013-02","No KG match: search proceeds without graph context","Negative","FR-013","P2",
"Query mentions unknown entities.",
"1. Ask: 'Weather forecast for Tokyo?'\n2. Check KG match.",
"'Tokyo' not in KG",
"No match. Empty graph context. Vector search proceeds. Answer from chunks only.",
"Vector search executed without graph augmentation.")

# FR-014
t("TC-FR014-01","Top-k retrieval + detailed answer with references","Happy Path","FR-014","P1",
"KG and vector index populated.",
"1. Ask 'What was APAC Q4 revenue?' (detailed).\n2. Verify chunks, answer, refs.",
"mode: detailed, k=10",
"Top-10 chunks by cosine similarity. Answer <=1000 words. Model: gemini-1.5-pro. Refs: [Source: Q4_Report.pdf, Page 12]. Latency <3s.",
"Answer with refs stored. Latency <3s.")

t("TC-FR014-02","Brief mode: <=150 words, fast model","Happy Path","FR-014","P1",
"Mode: Brief.",
"1. Toggle Brief.\n2. Ask 'Summarize Q4 performance.'",
"mode: brief",
"Answer <=150 words. Model: gemini-1.5-flash. Badge shows model.",
"Answer <=150 words.")

t("TC-FR014-03","No relevant chunks: fallback message","Negative","FR-014","P1",
"Unrelated question.",
"1. Ask 'Recipe for chocolate cake?'\n2. Check similarity.",
"No relevant docs in KB",
"Max similarity <0.5. Response: 'I couldn't find relevant information in the knowledge base to answer this question. Try rephrasing or uploading relevant documents.'",
"Fallback message. No fabrication.")

t("TC-FR014-04","P95 end-to-end latency under 3 seconds","Happy Path","FR-014","P1",
"Normal load.",
"1. Submit 100 questions.\n2. Record latencies.\n3. Calculate P95.",
"100 diverse questions",
"P95 latency <=3000ms.",
"Performance target met.")

t("TC-FR014-05","Max 10 references per answer","Boundary","FR-014","P2",
"Query matches 15+ documents.",
"1. Ask question matching many docs.\n2. Count references.",
"Query matching 15+ source docs",
"Max 10 references. Ranked by relevance (highest first). Each has doc_name, pages, score.",
"<=10 references.")

# FR-015
t("TC-FR015-01","Clickable reference opens document preview","Happy Path","FR-015","P1",
"Answer with references displayed.",
"1. View answer ref: Q4_Report.pdf, Page 12.\n2. Click reference.\n3. Verify preview.",
"Reference: Q4_Report.pdf, Page 12, score 0.94",
"Preview modal opens. Relevant chunk highlighted. Page 12 visible. Modal closeable.",
"Preview modal functional.")

t("TC-FR015-02","References stored as JSONB in Messages","Happy Path","FR-015","P2",
"Answer generated.",
"1. Query Messages table.\n2. Inspect references field.",
"Answer with 3 references",
"JSONB array: [{document_id, document_name, page_numbers, relevance_score}]. <=10 elements.",
"references JSONB properly structured.")

# ===== MODULE E: USER INTERFACE =====
# FR-016
t("TC-FR016-01","Drag-and-drop zone accepts files","Happy Path","FR-016","P1",
"On Upload page.",
"1. Drag report.pdf onto drop zone.\n2. Drop.",
"report.pdf, 2 MB",
"Zone highlights on drag-over. File accepted. Progress bar appears. Library shows Pending badge.",
"File uploaded.")

t("TC-FR016-02","Status badges: correct colors","Happy Path","FR-016","P2",
"Documents at all stages.",
"1. View library.\n2. Verify badge colors.",
"Docs: Pending, Processing, Completed, Failed",
"Pending=gray, Processing=blue animated, Completed=green, Failed=red with Retry button.",
"Correct badges.")

t("TC-FR016-03","Batch delete selected documents","Happy Path","FR-016","P2",
"5 documents. User has delete permission.",
"1. Select 3 checkboxes.\n2. Click Delete Selected.\n3. Confirm.",
"3 selected documents",
"Batch bar appears. Confirmation shown. 3 docs soft-deleted. Audit logged.",
"3 deleted. 2 remain.")

# FR-017
t("TC-FR017-01","Create conversation with auto-generated title","Happy Path","FR-017","P1",
"On main dashboard.",
"1. Click + New Conversation.\n2. Ask: 'What is Q4 APAC revenue?'\n3. Check sidebar.",
"First question: 'What is Q4 APAC revenue?'",
"Conversation created. Title auto-generated. Appears at top of sidebar. Truncated to 40 chars. msg_count=1.",
"Conversation record created.")

t("TC-FR017-02","Pin/unpin conversation","Happy Path","FR-017","P2",
"Existing conversations.",
"1. Pin a conversation.\n2. Verify pinned section.\n3. Unpin.",
"Conversation: 'Q4 Financial Performance'",
"Pinned: top of sidebar with pin icon, is_pinned=true. Unpinned: normal position, is_pinned=false.",
"Pin state toggled correctly.")

t("TC-FR017-03","Delete conversation with confirmation","Happy Path","FR-017","P2",
"Existing conversation.",
"1. Click delete.\n2. Read confirmation.\n3. Confirm.",
"Conversation with 3 messages",
"Dialog: 'Delete this conversation? This cannot be undone.' After confirm: soft-deleted, removed from sidebar.",
"is_deleted=true.")

t("TC-FR017-04","Search conversations by keyword","Happy Path","FR-017","P2",
"10+ conversations.",
"1. Type 'revenue' in search.\n2. Verify filtering.",
"3 of 10 conversations mention revenue",
"Only 3 matching shown. Case-insensitive. Updates as user types.",
"Filtered results correct.")

# FR-018
t("TC-FR018-01","Enter sends question; markdown answer rendered","Happy Path","FR-018","P1",
"Active conversation. KB populated.",
"1. Type question.\n2. Press Enter.\n3. View answer.",
"Question: 'List top 3 revenue regions.'",
"Enter submits. Loading dots shown. Markdown rendered (bold, lists, headings). Thumbs up/down buttons visible.",
"Answer displayed with markdown.")

t("TC-FR018-02","Shift+Enter adds newline","Happy Path","FR-018","P2",
"Active conversation.",
"1. Type 'Line one'.\n2. Shift+Enter.\n3. Type 'Line two'.\n4. Enter.",
"Multi-line input",
"New line added (textarea expands, max 5 lines). Enter sends full multi-line question.",
"Multi-line question submitted.")

t("TC-FR018-03","Rate answer with feedback","Happy Path","FR-018","P2",
"Answer displayed.",
"1. Click thumbs-up.\n2. Type feedback: 'Very accurate'.\n3. Submit.",
"Rating: thumbs up. Feedback: 'Very accurate'",
"user_rating stored. user_feedback='Very accurate'. Visual confirmation.",
"Rating and feedback in Messages.")

t("TC-FR018-04","Regenerate re-sends last question","Happy Path","FR-018","P2",
"Answer displayed.",
"1. Click Regenerate.\n2. Wait for new answer.",
"Original: 'What was Q4 revenue?'",
"Same question re-sent. New answer generated (may differ). Loading shown.",
"New answer generated.")

t("TC-FR018-05","Copy button copies to clipboard","Happy Path","FR-018","P3",
"Answer displayed.",
"1. Click Copy.\n2. Paste elsewhere.",
"Answer text about APAC revenue",
"Text copied to clipboard. Paste matches. Toast confirmation.",
"Clipboard contains answer.")

# FR-019
t("TC-FR019-01","Toggle Brief/Detailed modes","Happy Path","FR-019","P1",
"Default: Brief. Active conversation.",
"1. Verify Brief active (lightning icon).\n2. Click Detailed.\n3. Ask question.",
"Toggle Brief->Detailed. Question: 'Explain Q4 results.'",
"Toggle switches (book icon). Answer <=1000 words. Model: gemini-1.5-pro. Badge shown.",
"Detailed mode active.")

t("TC-FR019-02","Mode change does not affect previous answers","Happy Path","FR-019","P2",
"3 Brief answers exist.",
"1. Toggle to Detailed.\n2. Verify previous answers unchanged.\n3. Submit new question.",
"3 existing Brief answers",
"Previous 3 unchanged. New answer uses Detailed. No regeneration.",
"Previous answers intact.")

# FR-020
t("TC-FR020-01","Cached answer served <200ms with badge","Happy Path","FR-020","P1",
"Question previously cached. TTL valid.",
"1. Ask cached question.\n2. Measure response.\n3. Check badge.",
"Semantically identical question (cosine >=0.95)",
"Response <200ms. 'Cached' badge shown. is_cached=true. hit_count incremented.",
"Cache hit. hit_count++.")

t("TC-FR020-02","Regenerate bypasses cache","Happy Path","FR-020","P2",
"Cached answer shown.",
"1. Click Regenerate.\n2. Verify fresh generation.",
"Cached question",
"Fresh answer (2-3s). is_cached=false. May differ from cached version.",
"Fresh answer. is_cached=false.")

t("TC-FR020-03","Cache invalidated on new document ingestion","Happy Path","FR-020","P1",
"Cache populated. New document uploaded.",
"1. Verify cache entries.\n2. Upload new_report.pdf.\n3. Wait for ingestion.\n4. Ask cached question.",
"5 cache entries. New doc uploaded.",
"Cache invalidated. Fresh answer generated, may include new doc info.",
"Cache entries invalidated.")

t("TC-FR020-04","Brief/Detailed cached independently","Boundary","FR-020","P2",
"Same question in both modes.",
"1. Ask in Brief.\n2. Toggle Detailed, ask same.\n3. Toggle Brief, ask again.",
"Same question, both modes",
"Two separate cache entries. Step 3 serves Brief cache (<200ms). Different cache_keys.",
"Independent cache entries.")

t("TC-FR020-05","Cache TTL expiry after 7 days","Boundary","FR-020","P3",
"Cache entry from 8 days ago.",
"1. Query with expired cache.\n2. Verify miss.",
"Entry: expires_at = 7 days ago. Now: 8 days after creation.",
"Cache miss. Fresh answer. New cache entry created.",
"Expired cache bypassed.")

# FR-021
t("TC-FR021-01","First-login guided tour (5 steps)","Happy Path","FR-021","P2",
"New user first login.",
"1. Log in as new user.\n2. Step through tour.",
"new_analyst@adssoftek.com, first login",
"5 steps: Upload, Sidebar, Ask, Toggle, References. Each highlights UI element. Skippable. Does not repeat.",
"Tour completed. Not shown next login.")

t("TC-FR021-02","Keyboard shortcuts","Happy Path","FR-021","P2",
"Logged in on dashboard.",
"1. Ctrl+N.\n2. Ctrl+B.\n3. Ctrl+Enter.\n4. Esc.",
"Keyboard shortcuts",
"Ctrl+N: new conversation. Ctrl+B: toggle mode. Ctrl+Enter: send. Esc: close modal.",
"All shortcuts work.")

t("TC-FR021-03","Responsive design at 768px","Happy Path","FR-021","P2",
"Tablet width.",
"1. Resize to 768px.\n2. Navigate all screens.",
"Width: 768px",
"Sidebar collapses to overlay. No horizontal scroll. All features functional.",
"Responsive layout works.")

t("TC-FR021-04","WCAG 2.1 AA contrast compliance","Happy Path","FR-021","P2",
"Light and dark themes.",
"1. Check text contrast ratios.\n2. Verify focus indicators.\n3. Check ARIA labels.",
"Light: #1A1A2E on #FFFFFF. Dark: #FFFFFF on #1A1A2E.",
"Normal text >=4.5:1. Large text >=3:1. Focus indicators visible. ARIA labels on all controls.",
"WCAG AA met.")

# ===== MODULE F: ADMINISTRATION =====
# FR-022
t("TC-FR022-01","Admin creates new user","Happy Path","FR-022","P1",
"Admin logged in. User Management page.",
"1. Click Create User.\n2. Fill: email=analyst2@adssoftek.com, name=Priya Menon, role=analyst.\n3. Submit.",
"Email: analyst2@adssoftek.com, Name: Priya Menon, Role: analyst",
"User created. Temp password emailed. Status: active. Audit log: user.create. Welcome email sent.",
"User record in DB.")

t("TC-FR022-02","Deactivate user invalidates sessions","Happy Path","FR-022","P1",
"Target user active with session.",
"1. Edit analyst2@adssoftek.com.\n2. Set Inactive.\n3. Save.",
"analyst2@adssoftek.com, active, has session",
"Status=inactive. Sessions invalidated. Login disabled. Email sent. Audit logged.",
"User inactive. Sessions gone.")

t("TC-FR022-03","Non-admin blocked from user management","Permission","FR-022","P1",
"Logged in as Data Manager.",
"1. Access User Management.\n2. API: GET/POST /api/v1/users.",
"Role: Data Manager",
"Page not visible. API: 403 FORBIDDEN.",
"Access denied.")

t("TC-FR022-04","Soft-delete reassigns data to archive user","Happy Path","FR-022","P2",
"User has conversations and uploads.",
"1. Delete john@partner.com.\n2. Confirm.",
"john@partner.com with 5 conversations, 10 documents",
"Soft-deleted. Data reassigned to [archived] user. Removed from list. Cannot login.",
"is_deleted=true. Data reassigned.")

t("TC-FR022-05","Email immutable after creation","Negative","FR-022","P2",
"Editing existing user.",
"1. Edit user.\n2. Try changing email.\n3. Verify read-only.",
"Existing email. Attempted change.",
"Email field disabled. API: validation error on email change.",
"Email unchanged.")

# FR-023
t("TC-FR023-01","Update setting with audit log","Happy Path","FR-023","P1",
"Admin on Settings.",
"1. Change cache_ttl_hours: 168 -> 336.\n2. Save.",
"cache_ttl_hours: 168 -> 336",
"Updated. Audit: settings.update {old:168, new:336}. Notification to all admins.",
"Setting=336. Audit logged.")

t("TC-FR023-02","Reject invalid setting value","Negative","FR-023","P1",
"Admin on Settings.",
"1. Set cache_ttl_hours to -5.\n2. Save.",
"cache_ttl_hours: -5",
"Error: 'cache_ttl must be positive integer.' Setting unchanged.",
"Setting still 168.")

t("TC-FR023-03","Reset to Defaults","Happy Path","FR-023","P2",
"Settings modified.",
"1. Click Reset to Defaults.\n2. Confirm.",
"Modified: cache_ttl=336, chunking=fixed_size",
"All restored to factory defaults. Audit entries per setting.",
"All defaults restored.")

t("TC-FR023-04","Analyst blocked from settings","Permission","FR-023","P1",
"Logged in as Analyst.",
"1. Access Settings.\n2. API: GET /api/v1/admin/settings.",
"Role: Analyst",
"Not visible. API: 403 FORBIDDEN.",
"Access denied.")

# FR-024
t("TC-FR024-01","Real-time ingestion stats dashboard","Happy Path","FR-024","P1",
"Admin logged in. Various doc stages.",
"1. Admin > Ingestion Monitor.\n2. Verify stats.\n3. Check auto-refresh.",
"500 total, 10 pending, 5 processing, 480 completed, 5 failed",
"Cards show correct counts. Live job queue. Auto-refresh 30s. Manual refresh works.",
"Stats accurate.")

t("TC-FR024-02","Retry failed job from error table","Happy Path","FR-024","P1",
"Failed document exists.",
"1. Find failed doc in errors.\n2. Click Retry.",
"corrupt_report.pdf, failed, retry_count=3",
"Job re-queued. Status: queued -> processing.",
"Job retrying.")

t("TC-FR024-03","Data Manager can view ingestion logs","Permission","FR-024","P2",
"Logged in as Data Manager.",
"1. Access Ingestion Monitor.",
"Role: Data Manager",
"Full access to monitor, queue, errors. Retry available.",
"Access granted.")

# FR-025
t("TC-FR025-01","Search analytics charts and top questions","Happy Path","FR-025","P1",
"Admin logged in. Query history exists.",
"1. Admin > Analytics > Search.\n2. View charts.",
"1000 queries over 30 days",
"Query volume chart. Top 20 questions. Avg response time. Avg rating. Date filter works.",
"Analytics displayed.")

t("TC-FR025-02","CSV export of analytics","Happy Path","FR-025","P2",
"Admin on Analytics.",
"1. Click Export CSV.",
"User Activity, 30 days",
"CSV downloaded with correct headers/data matching chart.",
"CSV correct.")

t("TC-FR025-03","Analyst blocked from analytics","Permission","FR-025","P1",
"Logged in as Analyst.",
"1. Access Analytics.\n2. API: GET /api/v1/admin/analytics/overview.",
"Role: Analyst",
"Not visible. API: 403 FORBIDDEN.",
"Access denied.")

# ===== SECURITY & AUTH =====
t("TC-SEC-01","JWT access token expires after 1 hour","Integration","NFR-Security","P1",
"Authenticated user.",
"1. Login, get JWT.\n2. Wait 61 minutes.\n3. API call with expired token.",
"JWT TTL: 1 hour. Wait 61 min.",
"HTTP 401 UNAUTHORIZED. Must refresh or re-login.",
"Expired token rejected.")

t("TC-SEC-02","Refresh token renews access within 7 days","Integration","NFR-Security","P1",
"Access token expired. Refresh valid.",
"1. POST /api/v1/auth/refresh with valid refresh token.",
"Refresh token <7 days old",
"New access token issued, valid 1 hour.",
"Session continued.")

t("TC-SEC-03","Account locked after 5 failed logins","Negative","NFR-Security","P1",
"Active user.",
"1. Wrong password 5 times.\n2. Correct password on 6th.",
"analyst@adssoftek.com. Wrong: 'WrongPass!' x5",
"Locked 30 min after 5th fail. Email sent. 6th attempt rejected. Auto-unlocks after 30 min.",
"Account locked. Notification sent.")

t("TC-SEC-04","Password policy enforcement","Negative","NFR-Security","P1",
"User setting password.",
"1. Try 'short' (too short).\n2. 'alllower1!' (no uppercase).\n3. 'NoNumbers!!' (no digit).\n4. 'NoSpecial1A' (no special).\n5. 'SecureP@ss1' (valid).",
"5 password attempts",
"1-4 rejected with specific errors. 5: accepted. Min 8 chars, 1 upper, 1 number, 1 special.",
"Only valid password accepted.")

t("TC-SEC-05","Rate limiting: 100/min general, 20/min LLM","Boundary","NFR-Security","P1",
"Authenticated.",
"1. 101 general requests/min.\n2. 21 LLM requests/min.",
"101 GET, 21 POST requests",
"101st general: 429. 21st LLM: 429. Headers: X-RateLimit-Limit, Remaining, Reset.",
"Rate limits enforced.")

t("TC-SEC-06","Concurrent session limit of 5","Boundary","NFR-Security","P2",
"User logs in from multiple devices.",
"1. Login from 5 devices.\n2. Attempt 6th.",
"5 active sessions",
"6th rejected or oldest invalidated. Max 5 enforced.",
"5 session limit enforced.")

t("TC-SEC-07","Sessions revoked on password change","Integration","NFR-Security","P2",
"3 active sessions.",
"1. Change password on device 1.\n2. API call from devices 2,3.",
"3 sessions. Password changed on session 1.",
"Devices 2,3: 401. Must re-auth.",
"Other sessions invalidated.")

# ===== WORKFLOW STATE TRANSITIONS =====
t("TC-WF-01","Full document ingestion happy path","Workflow","FR-001 to FR-011","P1",
"User uploads native PDF.",
"1. Upload Q4_Report.pdf.\n2. Monitor status transitions.",
"Q4_Report.pdf, 42 pages, native text, 2.4 MB",
"Pending -> Processing -> Converting -> OCR_Check (skip OCR) -> Chunking -> Vectorizing -> Entity_Extraction -> Graph_Building -> Completed. All side effects at each step.",
"ingestion_status=completed.")

t("TC-WF-02","Ingestion failure and retry workflow","Workflow","FR-001,FR-002","P1",
"Conversion fails.",
"1. Upload problematic.xlsx.\n2. Observe retries.",
"problematic.xlsx causing conversion error",
"Converting -> Error -> Retrying (x3). Success: continues. All fail: status=failed, user notified.",
"Either completed or failed.")

t("TC-WF-03","Query workflow: cache miss path","Workflow","FR-012 to FR-015","P1",
"Active conversation. No cache match.",
"1. Ask question.\n2. Monitor workflow.",
"'What was APAC Q4 revenue?' detailed mode",
"Processing_Query -> Expanding_Intent -> Cache_Check (miss) -> Graph_Query -> Vector_Search -> Generating_Answer -> Active. Answer stored.",
"Full query workflow completed.")

t("TC-WF-04","Query workflow: cache hit shortcut","Workflow","FR-020","P1",
"Cached question.",
"1. Ask cached question.\n2. Monitor shortcut.",
"Same question, cosine >=0.95",
"Processing_Query -> Expanding_Intent -> Cache_Hit -> Active. Skips graph/vector/LLM. <200ms.",
"Cache hit path.")

t("TC-WF-05","User lifecycle: create to delete","Workflow","FR-022","P1",
"Admin manages user.",
"1. Create -> Login -> 5 fails -> Lock -> Unlock -> Deactivate -> Reactivate -> Delete.",
"test@adssoftek.com through full lifecycle",
"New->Active->Active->Locked(30min)->Active->Inactive->Active->Deleted. All side effects.",
"Full lifecycle traversed.")

# ===== INTEGRATION =====
t("TC-INT-01","E2E: Upload to answer with reference","Integration","FR-001 to FR-015","P1",
"Clean system. Analyst logged in.",
"1. Upload Q4_Report.pdf.\n2. Wait for full ingestion.\n3. Create conversation.\n4. Ask about APAC Q4 revenue.\n5. Verify answer cites document.",
"Q4_Report.pdf with 'APAC revenue $1.2B'. Detailed mode.",
"Full pipeline: upload -> ingest -> chunk -> vectorize -> KG -> query -> answer. Answer references Q4_Report.pdf with pages.",
"End-to-end works.")

t("TC-INT-02","Upload + KG stats visible in admin","Integration","FR-001,FR-009,FR-025","P1",
"Admin. Clean system.",
"1. Upload company_profile.pdf.\n2. Wait ingestion.\n3. Admin > KG stats.",
"Doc mentioning ADS Softek, Rajesh Kumar, APAC",
"KG stats show new nodes/edges. Entities searchable via API.",
"Stats reflect new entities.")

t("TC-INT-03","Multi-user concurrent queries","Integration","FR-014,FR-017","P2",
"Two Analysts. Same KB.",
"1. Analyst A asks question.\n2. Analyst B asks same simultaneously.",
"Identical questions within 1 second",
"Both get correct answers. No cross-contamination. Second may hit cache.",
"Isolation maintained.")

t("TC-INT-04","Graceful degradation: LLM unavailable","Integration","NFR-Availability","P1",
"LLM service down.",
"1. Simulate LLM 503.\n2. Submit question.",
"LLM returns 503. Vector search works.",
"Raw chunk matches returned without synthesis. Error message about reduced functionality. No 500 error.",
"Graceful degradation.")

t("TC-INT-05","Graceful degradation: OCR unavailable","Integration","NFR-Availability","P2",
"Document AI down.",
"1. Upload scanned_doc.jpeg when OCR down.",
"OCR service 503",
"Document queued as pending_ocr. File stored (not lost). Resumes when service recovers.",
"Queued, not failed.")

# ===== NOTIFICATIONS =====
t("TC-NOTIF-01","Welcome email on user creation","Integration","FR-022,Notifications","P2",
"Admin creates user.",
"1. Create user: new@adssoftek.com.\n2. Check email.",
"new@adssoftek.com",
"Welcome email with temp password and login instructions.",
"Email delivered.")

t("TC-NOTIF-02","Ingestion failure notifies uploader + data managers","Integration","FR-002,Notifications","P1",
"Document fails after 3 retries.",
"1. Upload failing doc.\n2. Wait for failure.\n3. Check notifications.",
"failing_doc.pdf. Uploader: analyst@adssoftek.com",
"In-app + email to uploader. In-app to Data Managers. Critical alert.",
"Notifications sent.")

t("TC-NOTIF-03","Setting change notifies all admins","Integration","FR-023,Notifications","P2",
"Admin changes setting.",
"1. Change cache_ttl.\n2. Check other admins' notifications.",
"2 admins in system",
"In-app to all admins: '[Setting] changed from [old] to [new] by [admin].'",
"All admins notified.")

# ===== PERFORMANCE =====
t("TC-PERF-01","Page load <2s, SPA nav <500ms","Integration","NFR-Performance","P1",
"Normal load.",
"1. Full page load.\n2. SPA navigation.",
"Chrome 90+, broadband",
"Initial load <2s. Subsequent nav <500ms. Skeleton loaders shown.",
"Performance met.")

t("TC-PERF-02","500 concurrent users","Integration","NFR-Performance","P1",
"Load test environment.",
"1. Simulate 500 users.\n2. Mix: 40% search, 30% browse, 20% upload, 10% admin.\n3. Measure.",
"500 VUs, 15 min, 5 min ramp",
"P95 response <=3s. No 5xx errors. Auto-scaling triggers if needed.",
"SLAs met at 500 users.")

t("TC-PERF-03","HNSW vector search <100ms","Integration","FR-008,NFR-Performance","P2",
"100K+ embeddings indexed.",
"1. Nearest-neighbor search.\n2. Measure time.",
"768-dim query vector, 100K+ index",
"<100ms. Results ranked by cosine similarity. HNSW used.",
"Sub-100ms confirmed.")

# ===== DATA VALIDATION =====
t("TC-DATA-01","Unique RFC 5322 email validation","Negative","FR-022,Data Model","P1",
"Admin creating users.",
"1. Duplicate email.\n2. Invalid format.\n3. Valid unique.",
"'admin@adssoftek.com' (dup), 'not-an-email' (bad), 'valid@adssoftek.com'",
"Dup: 409 CONFLICT. Invalid: 422 VALIDATION_ERROR. Valid: created.",
"Only valid unique email accepted.")

t("TC-DATA-02","file_size_bytes: >0 and <=104857600","Boundary","FR-001,Data Model","P2",
"API validation.",
"1. 0-byte file.\n2. 104,857,600 bytes.\n3. 104,857,601 bytes.",
"0 bytes, 100 MB, 100 MB + 1 byte",
"0-byte rejected. 100 MB accepted. 100 MB+1 rejected.",
"Boundary enforced.")

t("TC-DATA-03","Conversation title max 300 chars","Boundary","FR-017,Data Model","P2",
"Creating/renaming conversation.",
"1. 300-char title.\n2. 301-char title.",
"300 chars (valid), 301 chars (invalid)",
"300 accepted. 301 rejected/truncated. Sidebar shows 40-char truncation.",
"Title length validated.")

t("TC-DATA-04","Message content max 50,000 chars","Boundary","FR-018,Data Model","P3",
"Very long LLM answer.",
"1. Generate maximum-length answer.\n2. Check DB content length.",
"Question triggering long answer",
"Content <=50,000 chars. Truncated gracefully if exceeded.",
"Within limit.")

t("TC-DATA-05","User rating 1-5 range","Boundary","FR-018,Data Model","P2",
"Rating an answer.",
"1. Rate 0.\n2. Rate 6.\n3. Rate 3.",
"Ratings: 0, 6, 3",
"0 rejected. 6 rejected. 3 accepted.",
"Only 1-5 accepted.")

t("TC-DATA-06","Max 3 automatic retries for ingestion jobs","Boundary","FR-002,Data Model","P2",
"Failing ingestion job.",
"1. Job fails repeatedly.\n2. Verify count.\n3. No 4th auto retry.",
"Consistently failing job",
"attempt_count: 1,2,3. After 3: status=failed. No auto 4th. Manual retry via admin OK.",
"Max 3 retries.")

t("TC-DATA-07","KG confidence 0.0-1.0; threshold 0.75","Boundary","FR-009,Data Model","P2",
"Entity extraction.",
"1. Score 0.74 (below).\n2. Score 0.75 (boundary).\n3. Score 1.0.",
"Confidence: 0.74, 0.75, 1.0",
"0.74 discarded. 0.75 stored. 1.0 stored.",
"Threshold enforced.")

t("TC-DATA-08","Audit logs append-only immutable","Integration","Data Model,NFR-Security","P1",
"Audit log exists.",
"1. Create entry.\n2. Try UPDATE.\n3. Try DELETE.",
"Audit entry: document.upload",
"UPDATE fails/prohibited. DELETE fails/prohibited. created_at immutable.",
"Logs cannot be modified.")

# ===== API CONTRACT =====
t("TC-API-01","Standardized error response format","Integration","API Requirements","P1",
"Various error scenarios.",
"1. Invalid body -> 422.\n2. No auth -> 401.\n3. Forbidden -> 403.\n4. Not found -> 404.",
"Invalid body, no auth, wrong role, bad ID",
"All errors: {\"error\":{\"code\":\"...\",\"message\":\"...\",\"field\":\"...\",\"details\":{}}}. 422=VALIDATION_ERROR. 401=UNAUTHORIZED. 403=FORBIDDEN. 404=NOT_FOUND.",
"Consistent format.")

t("TC-API-02","Rate limit headers in all responses","Integration","API Requirements","P2",
"Authenticated request.",
"1. Any API request.\n2. Check headers.",
"GET /api/v1/documents",
"Headers: X-RateLimit-Limit, X-RateLimit-Remaining, X-RateLimit-Reset.",
"Headers present.")


# ===== BUILD DOCUMENT =====
def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(10)
    for lv in range(1,4):
        doc.styles[f"Heading {lv}"].font.name = "Arial"
        doc.styles[f"Heading {lv}"].font.color.rgb = RGBColor.from_string("003366")

    # Title page
    for _ in range(6): doc.add_paragraph("")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IntelliRAG"); r.font.size = Pt(36); r.bold = True
    r.font.color.rgb = RGBColor.from_string("003366"); r.font.name = "Arial"
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Test Case Specification v1.0"); r2.font.size = Pt(24)
    r2.font.color.rgb = RGBColor.from_string("2E75B6"); r2.font.name = "Arial"
    doc.add_paragraph("")
    for line in ["Date: March 17, 2026","Prepared for: ADS Softek","Status: Draft"]:
        px = doc.add_paragraph(); px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rx = px.add_run(line); rx.font.size = Pt(12); rx.font.name = "Arial"
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for item in ["1. Test Coverage Summary","2. Traceability Matrix",
                  "3. Test Cases by Module","   3.1 Module A: File Ingestion",
                  "   3.2 Module B: Storage, Chunking & Vectorization",
                  "   3.3 Module C: Knowledge Graph","   3.4 Module D: RAG Search",
                  "   3.5 Module E: User Interface","   3.6 Module F: Administration",
                  "   3.7 Authentication & Security","   3.8 Workflow State Transitions",
                  "   3.9 Cross-Module Integration","   3.10 Notifications",
                  "   3.11 Performance","   3.12 Data Validation","   3.13 API Contract"]:
        doc.add_paragraph(item)
    doc.add_page_break()

    # 1. Coverage Summary
    total = len(TCS)
    cats = {}; pris = {}
    for tc in TCS:
        cats[tc[2]] = cats.get(tc[2],0)+1
        pris[tc[4]] = pris.get(tc[4],0)+1

    doc.add_heading("1. Test Coverage Summary", level=1)
    p = doc.add_paragraph(); r = p.add_run(f"Total Test Cases: {total}")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Arial"

    doc.add_heading("By Category", level=2)
    tb = doc.add_table(rows=1+len(cats), cols=2); tb.style = "Table Grid"
    hdr_row(tb, ["Category","Count"])
    for i,(c,n) in enumerate(sorted(cats.items()),1):
        cell_txt(tb.rows[i].cells[0], c); cell_txt(tb.rows[i].cells[1], str(n))

    doc.add_heading("By Priority", level=2)
    tb2 = doc.add_table(rows=1+len(pris), cols=2); tb2.style = "Table Grid"
    hdr_row(tb2, ["Priority","Count"])
    for i,(p,n) in enumerate(sorted(pris.items()),1):
        cell_txt(tb2.rows[i].cells[0], p); cell_txt(tb2.rows[i].cells[1], str(n))

    modules_map = {
        "Module A: File Ingestion (FR-001 to FR-004)": ["FR-001","FR-002","FR-003","FR-004"],
        "Module B: Storage & Vectorization (FR-005 to FR-008)": ["FR-005","FR-006","FR-007","FR-008"],
        "Module C: Knowledge Graph (FR-009 to FR-011)": ["FR-009","FR-010","FR-011"],
        "Module D: RAG Search (FR-012 to FR-015)": ["FR-012","FR-013","FR-014","FR-015"],
        "Module E: User Interface (FR-016 to FR-021)": ["FR-016","FR-017","FR-018","FR-019","FR-020","FR-021"],
        "Module F: Administration (FR-022 to FR-025)": ["FR-022","FR-023","FR-024","FR-025"],
    }
    doc.add_heading("By Module", level=2)
    tb3 = doc.add_table(rows=1+len(modules_map), cols=2); tb3.style = "Table Grid"
    hdr_row(tb3, ["Module","Count"])
    for i,(m,frs) in enumerate(modules_map.items(),1):
        cnt = sum(1 for tc in TCS if any(f in tc[3] for f in frs))
        cell_txt(tb3.rows[i].cells[0], m); cell_txt(tb3.rows[i].cells[1], str(cnt))
    doc.add_page_break()

    # 2. Traceability Matrix
    doc.add_heading("2. Traceability Matrix", level=1)
    doc.add_paragraph("Maps each Functional Requirement to its test case IDs.")
    all_frs = [f"FR-{str(i).zfill(3)}" for i in range(1,26)]
    tb4 = doc.add_table(rows=1+len(all_frs), cols=3); tb4.style = "Table Grid"
    hdr_row(tb4, ["FR ID","Test Case IDs","Count"])
    for i,fr in enumerate(all_frs,1):
        ids = [tc[0] for tc in TCS if fr in tc[3]]
        cell_txt(tb4.rows[i].cells[0], fr, bold=True, sz=8)
        cell_txt(tb4.rows[i].cells[1], ", ".join(ids) if ids else "See integration/workflow tests", sz=8)
        cell_txt(tb4.rows[i].cells[2], str(len(ids)), sz=8)
    doc.add_page_break()

    # 3. Test Cases
    doc.add_heading("3. Test Cases by Module", level=1)
    sections = [
        ("3.1 Module A: File Ingestion (FR-001 to FR-004)", ["TC-FR001","TC-FR002","TC-FR003","TC-FR004"]),
        ("3.2 Module B: Storage, Chunking & Vectorization (FR-005 to FR-008)", ["TC-FR005","TC-FR006","TC-FR007","TC-FR008"]),
        ("3.3 Module C: Knowledge Graph (FR-009 to FR-011)", ["TC-FR009","TC-FR010","TC-FR011"]),
        ("3.4 Module D: RAG Search (FR-012 to FR-015)", ["TC-FR012","TC-FR013","TC-FR014","TC-FR015"]),
        ("3.5 Module E: User Interface (FR-016 to FR-021)", ["TC-FR016","TC-FR017","TC-FR018","TC-FR019","TC-FR020","TC-FR021"]),
        ("3.6 Module F: Administration (FR-022 to FR-025)", ["TC-FR022","TC-FR023","TC-FR024","TC-FR025"]),
        ("3.7 Authentication & Security", ["TC-SEC"]),
        ("3.8 Workflow State Transitions", ["TC-WF"]),
        ("3.9 Cross-Module Integration", ["TC-INT"]),
        ("3.10 Notifications", ["TC-NOTIF"]),
        ("3.11 Performance", ["TC-PERF"]),
        ("3.12 Data Validation", ["TC-DATA"]),
        ("3.13 API Contract", ["TC-API"]),
    ]
    for title, prefixes in sections:
        doc.add_heading(title, level=2)
        matched = [tc for tc in TCS if any(tc[0].startswith(p) for p in prefixes)]
        for tc in matched:
            add_tc(doc, tc)
        doc.add_page_break()

    out = "/Users/n15318/RAG-app/docs/IntelliRAG_TestCases_v1.0.docx"
    doc.save(out)
    print(f"SUCCESS: Generated {out}")
    print(f"Total: {total} test cases")
    print(f"Categories: {dict(sorted(cats.items()))}")
    print(f"Priorities: {dict(sorted(pris.items()))}")

if __name__ == "__main__":
    build()
