"""Normalizes documents to text — PDF, DOCX, XLSX, XLS, DOC, images, with OCR fallback."""

import logging
import os
import unicodedata
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from ..db import get_connection, get_cursor
from ..storage import storage_client
from ..config import config
from .ocr_provider import (
    _document_ai_enabled,
    ocr_image,
    ocr_pdf,
    ocr_pdf_pages,
    reset_document_ai_client,
)

logger = logging.getLogger(__name__)

IMAGE_MIME_TYPES = {
    "image/jpeg", "image/png", "image/tiff", "image/bmp", "image/gif", "image/webp",
}
MIN_USABLE_TEXT_CHARS = 10


def _has_usable_text(text: str) -> bool:
    return bool(text and len(text.strip()) >= MIN_USABLE_TEXT_CHARS)


def _retry_full_pdf_ocr(local_path: str, document_id: str, text: str, page_count: int) -> tuple[str, int]:
    """Retry full PDF OCR once with a fresh Document AI client before failing."""
    if _has_usable_text(text):
        return text, page_count

    logger.info(
        "PDF text extraction yielded minimal content, trying full OCR",
        extra={"document_id": document_id},
    )
    text, page_count = ocr_pdf(local_path, document_id)
    if _has_usable_text(text) or not _document_ai_enabled():
        return text, page_count

    logger.warning(
        "Full OCR returned minimal content, retrying with a fresh Document AI client",
        extra={"document_id": document_id},
    )
    reset_document_ai_client()
    return ocr_pdf(local_path, document_id)


def normalize_document(document_id: str, workspace_id: str):
    """Extract text from document based on MIME type."""
    with get_connection() as conn:
        with get_cursor(conn) as cur:
            cur.execute(
                "SELECT file_path, mime_type, file_name FROM document WHERE document_id = %s",
                (document_id,)
            )
            doc = cur.fetchone()
            if not doc:
                raise ValueError(f"Document {document_id} not found")

    file_path = doc["file_path"]
    mime_type = doc["mime_type"]

    # Download from GCS if needed
    local_path = storage_client.download_to_temp(file_path)

    try:
        if mime_type == "application/pdf":
            text, page_count = _extract_pdf(local_path, document_id)
        elif mime_type == "application/msword":
            text, page_count = _extract_doc(local_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text, page_count = _extract_docx(local_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            text, page_count = _extract_xlsx(local_path)
        elif mime_type == "application/vnd.ms-excel":
            text, page_count = _extract_xls(local_path)
        elif mime_type in ("text/plain", "text/csv"):
            with open(local_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            page_count = 1
        elif mime_type == "text/markdown":
            with open(local_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            text = _strip_markdown(text)
            page_count = 1
        elif mime_type in IMAGE_MIME_TYPES:
            text, page_count = ocr_image(local_path, document_id, mime_type)
        else:
            raise ValueError(f"Unsupported MIME type for normalization: {mime_type}")

        if not _has_usable_text(text):
            if mime_type == "application/pdf":
                text, page_count = _retry_full_pdf_ocr(local_path, document_id, text, page_count)
            if not _has_usable_text(text):
                raise ValueError("Document normalization produced no usable text")

        # Unicode NFC normalization (FR-004/AC-04)
        text = unicodedata.normalize("NFC", text)

        # Extract metadata
        extracted_meta = _extract_file_metadata(local_path, mime_type)
        language = _detect_language(text)

        # FR-007: Inject page annotations into text for downstream chunker
        if page_count > 1 and "\f" in text:
            pages = text.split("\f")
            annotated_pages = []
            for idx, page_text in enumerate(pages):
                annotated_pages.append(f"<!-- PAGE {idx + 1} -->\n{page_text}")
            text = "\f".join(annotated_pages)

        # Store extraction results
        with get_connection() as conn:
            with get_cursor(conn) as cur:
                cur.execute("""
                    INSERT INTO extraction_result (document_id, extraction_type, content, metadata)
                    VALUES (%s, 'TEXT', %s, %s)
                """, (document_id, text, f'{{"page_count": {page_count}}}'))

                # FR-007: Flag low-confidence OCR documents for review
                review_required = False
                if hasattr(normalize_document, '_avg_confidence'):
                    avg_conf = normalize_document._avg_confidence
                    if avg_conf < config.OCR_REVIEW_THRESHOLD:
                        review_required = True
                        logger.warning("Low OCR confidence, flagging for review",
                                       extra={"document_id": document_id, "avg_confidence": avg_conf})

                cur.execute("""
                    UPDATE document
                    SET page_count = %s, extracted_metadata = %s, language = %s,
                        review_required = %s, updated_at = now()
                    WHERE document_id = %s
                """, (page_count, _json_dumps(extracted_meta), language, review_required, document_id))

        logger.info(f"Document {document_id} normalized: {len(text)} chars, {page_count} pages, lang={language}")
    finally:
        storage_client.cleanup_temp(local_path, file_path)


def _extract_file_metadata(file_path: str, mime_type: str) -> dict:
    """Extract metadata from PDF or DOCX files."""
    meta = {}
    try:
        if mime_type == "application/pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                if pdf.metadata:
                    for key in ("Title", "Author", "Subject", "Creator", "Producer", "CreationDate", "Keywords"):
                        if pdf.metadata.get(key):
                            meta[key.lower()] = str(pdf.metadata[key])
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            props = doc.core_properties
            if props.title:
                meta["title"] = props.title
            if props.author:
                meta["author"] = props.author
            if props.subject:
                meta["subject"] = props.subject
            if props.created:
                meta["created"] = str(props.created)
            if props.modified:
                meta["modified"] = str(props.modified)
    except Exception as e:
        logger.warning(f"Metadata extraction failed: {e}")
    return meta


def _detect_language(text: str) -> str:
    """Detect document language using langdetect."""
    try:
        from langdetect import detect
        # Use first 5000 chars for speed
        sample = text[:5000]
        return detect(sample)
    except Exception:
        return "en"


def _strip_markdown(text: str) -> str:
    """Strip markdown syntax for plain text processing (FR-002/AC-03)."""
    import re
    # Remove headers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove bold/italic
    text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
    text = re.sub(r'_{1,3}([^_]+)_{1,3}', r'\1', text)
    # Remove links [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # Remove images ![alt](url)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    # Remove code blocks
    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove horizontal rules
    text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    return text.strip()


def _extract_doc(file_path: str) -> tuple:
    """Extract text from legacy .doc files: convert to DOCX via LibreOffice, then parse (FR-002/AC-02)."""
    import subprocess
    import shutil
    import tempfile

    # Try LibreOffice: convert DOC -> DOCX, then process through _extract_docx (FR-002/AC-02)
    if shutil.which("libreoffice") or shutil.which("soffice"):
        cmd = "libreoffice" if shutil.which("libreoffice") else "soffice"
        with tempfile.TemporaryDirectory() as tmpdir:
            try:
                subprocess.run(
                    [cmd, "--headless", "--convert-to", "docx", "--outdir", tmpdir, file_path],
                    check=True, capture_output=True, timeout=30,
                )
                docx_files = [f for f in os.listdir(tmpdir) if f.endswith(".docx")]
                if docx_files:
                    docx_path = os.path.join(tmpdir, docx_files[0])
                    return _extract_docx(docx_path)
            except Exception as e:
                logger.warning("LibreOffice DOC->DOCX conversion failed, trying fallbacks",
                               extra={"error": str(e)})

    # Try textutil (macOS built-in)
    if shutil.which("textutil"):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            subprocess.run(
                ["textutil", "-convert", "txt", "-output", tmp_path, file_path],
                check=True, capture_output=True, timeout=60,
            )
            with open(tmp_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            return text, 1
        except Exception as e:
            logger.warning("textutil failed for .doc", extra={"error": str(e)})
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # Try antiword
    if shutil.which("antiword"):
        try:
            result = subprocess.run(
                ["antiword", file_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout, 1
        except Exception as e:
            logger.warning("antiword failed for .doc", extra={"error": str(e)})

    raise ValueError("No .doc converter available (install libreoffice, textutil, or antiword)")


def _estimate_native_confidence(page) -> float:
    """Estimate confidence of native PDF text extraction (FR-007).

    Heuristic: ratio of printable chars to total chars, penalizing
    very short text, excessive whitespace, or encoding artifacts.
    """
    text = page.extract_text() or ""
    if len(text.strip()) < 10:
        return 0.0
    total = len(text)
    printable = sum(1 for c in text if c.isprintable() or c in ('\n', '\t'))
    ratio = printable / total if total > 0 else 0.0
    # Penalize very short pages (likely scanned)
    length_factor = min(1.0, len(text.strip()) / 200)
    return ratio * length_factor


def _extract_pdf(file_path: str, document_id: str = "") -> tuple:
    """Extract text from PDF with confidence-based OCR fallback (FR-003/AC-02, FR-007)."""
    import pdfplumber
    texts = []
    ocr_page_indices = []
    page_confidences = []
    confidence_threshold = config.OCR_NATIVE_CONFIDENCE_THRESHOLD

    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        for page_idx, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    rows = []
                    for row in table:
                        cells = [str(cell) if cell else "" for cell in row]
                        rows.append(" | ".join(cells))
                    text += "\n\n<!-- TABLE_START -->\n" + "\n".join(rows) + "\n<!-- TABLE_END -->"

            # FR-007: Confidence-based OCR trigger (not just text < 10 chars)
            native_conf = _estimate_native_confidence(page)
            page_confidences.append(native_conf)

            if native_conf < confidence_threshold:
                ocr_page_indices.append(page_idx)
                logger.info("Page below confidence threshold, marking for OCR",
                            extra={"document_id": document_id, "page": page_idx + 1,
                                   "confidence": f"{native_conf:.2f}", "threshold": confidence_threshold})
            texts.append(text)

    # OCR only the pages that need it (FR-003/AC-02)
    if ocr_page_indices:
        logger.info("Running per-page OCR on low-confidence pages",
                     extra={"document_id": document_id, "page_count": len(ocr_page_indices)})
        ocr_results = ocr_pdf_pages(file_path, document_id, ocr_page_indices)
        for page_idx, ocr_text in ocr_results.items():
            if ocr_text and len(ocr_text.strip()) >= 10:
                texts[page_idx] = ocr_text

    # Store average confidence for review flagging
    avg_conf = sum(page_confidences) / len(page_confidences) if page_confidences else 1.0
    normalize_document._avg_confidence = avg_conf

    return "\f".join(texts), page_count


def _extract_docx(file_path: str) -> tuple:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        paragraphs.append("<!-- TABLE_START -->\n" + "\n".join(table_rows) + "\n<!-- TABLE_END -->")
    return "\n\n".join(paragraphs), 1


def _extract_xlsx(file_path: str) -> tuple:
    """Extract XLSX to CSV-like text with per-sheet separation and timeout (FR-002/AC-01, BR-01)."""

    def _do_extract(path: str) -> tuple:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        sheet_count = len(wb.worksheets)
        texts = []
        for sheet in wb.worksheets:
            rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) if cell is not None else "" for cell in row]
                if any(cells):
                    rows.append(",".join(cells))
            if rows:
                # Each sheet clearly separated with header (FR-002/BR-01)
                csv_content = "\n".join(rows)
                texts.append(f"Sheet: {sheet.title}\n{csv_content}")
        wb.close()
        return "\n\n".join(texts), sheet_count

    # Run with timeout using ThreadPoolExecutor (FR-002/AC-01)
    timeout_s = config.XLSX_TIMEOUT_S
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(_do_extract, file_path)
        try:
            return future.result(timeout=timeout_s)
        except FuturesTimeoutError:
            logger.error("XLSX extraction timed out",
                         extra={"timeout_s": timeout_s, "file_path": file_path})
            future.cancel()
            raise ValueError(f"XLSX extraction timed out after {timeout_s}s")


def _extract_xls(file_path: str) -> tuple:
    """Extract legacy .xls files using xlrd (FR-002/AC-01)."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_path)
        texts = []
        for sheet in wb.sheets():
            rows = []
            for row_idx in range(sheet.nrows):
                cells = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                if any(cells):
                    rows.append(",".join(cells))
            if rows:
                texts.append(f"Sheet: {sheet.name}\n" + "\n".join(rows))
        return "\n\n".join(texts), len(wb.sheets())
    except ImportError:
        raise ValueError("xlrd is required for .xls support. Install with: pip install xlrd")


def _json_dumps(obj: dict) -> str:
    """Safe JSON serialization."""
    import json
    try:
        return json.dumps(obj, default=str)
    except Exception:
        return "{}"
